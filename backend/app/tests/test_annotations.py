import uuid
import pytest
from app.core.database import (
    init_db, create_project, create_document,
    add_annotation, get_annotations, update_annotation, delete_annotation
)
from app.api.annotations import api_get_annotations, api_create_annotation, CreateAnnotationBody, api_update_annotation, UpdateAnnotationBody, api_delete_annotation

@pytest.fixture(autouse=True)
def setup_db():
    init_db()

def test_annotation_crud_database():
    pid = f"proj_test_{uuid.uuid4().hex[:8]}"
    did = f"doc_test_{uuid.uuid4().hex[:8]}"
    proj = create_project(pid, "test_annot_proj")
    doc = create_document(did, proj["id"], "test.docx", "test.docx")
    
    # 1. Add annotation
    a1 = add_annotation(
        document_id=doc["id"],
        paragraph_idx=0,
        selected_text="西瓜少年",
        content="主人公名字与回忆象征",
        paragraph_uuid="p_uuid_0"
    )
    assert a1["id"].startswith("ann_")
    assert a1["selected_text"] == "西瓜少年"
    assert a1["content"] == "主人公名字与回忆象征"
    
    # 2. Get annotations
    annots = get_annotations(doc["id"])
    assert len(annots) == 1
    assert annots[0]["id"] == a1["id"]
    
    # 3. Update annotation
    updated = update_annotation(a1["id"], "更新后的注释文本")
    assert updated is not None
    assert updated["content"] == "更新后的注释文本"
    
    # 4. Delete annotation
    res_del = delete_annotation(a1["id"])
    assert res_del is True
    assert len(get_annotations(doc["id"])) == 0

@pytest.mark.anyio
async def test_annotation_api():
    pid = f"proj_test_{uuid.uuid4().hex[:8]}"
    did = f"doc_test_{uuid.uuid4().hex[:8]}"
    proj = create_project(pid, "test_api_annot_proj")
    doc = create_document(did, proj["id"], "test.docx", "test.docx")
    
    # 1. POST API
    body = CreateAnnotationBody(
        paragraph_idx=1,
        selected_text="冰镇汽水",
        content="童年夏天消暑饮品",
        paragraph_uuid="p_uuid_1"
    )
    res_create = await api_create_annotation(proj["id"], body)
    assert res_create["status"] == "ok"
    ann_id = res_create["annotation"]["id"]
    
    # 2. GET API
    res_get = await api_get_annotations(proj["id"])
    assert len(res_get) == 1
    assert res_get[0]["id"] == ann_id
    
    # 3. PUT API
    res_put = await api_update_annotation(proj["id"], ann_id, UpdateAnnotationBody(content="修改后的汽水说明"))
    assert res_put["status"] == "ok"
    assert res_put["annotation"]["content"] == "修改后的汽水说明"
    
    # 4. DELETE API
    res_del = await api_delete_annotation(proj["id"], ann_id)
    assert res_del["status"] == "ok"
    
    res_get_after = await api_get_annotations(proj["id"])
    assert len(res_get_after) == 0

@pytest.mark.anyio
async def test_export_comment_mode_native_openxml(tmp_path):
    from docx import Document
    from app.core.database import insert_paragraphs
    from app.api.apply import export_document
    
    pid = f"proj_exp_{uuid.uuid4().hex[:8]}"
    did = f"doc_exp_{uuid.uuid4().hex[:8]}"
    
    dummy_docx = tmp_path / "dummy.docx"
    doc_init = Document()
    doc_init.add_paragraph("第一段：西瓜少年的夏日故事。")
    doc_init.save(str(dummy_docx))
    
    proj = create_project(pid, "导出批注测试")
    doc = create_document(did, proj["id"], "dummy.docx", str(dummy_docx))
    insert_paragraphs(did, [(0, "第一段：西瓜少年的夏日故事。", "Normal")])
    
    add_annotation(
        document_id=did,
        paragraph_idx=0,
        selected_text="西瓜少年",
        content="主角称号说明"
    )
    
    res = await export_document(pid, export_mode="comment")
    assert hasattr(res, "path")
    out_docx = Document(res.path)
    
    # 验证导出的 docx 包含 word/comments.xml relationship 并且 XML 内存在批注
    comments_found = False
    for rel in out_docx.part.rels.values():
        if "comments" in str(rel.target_ref):
            comments_found = True
            blob_str = rel.target_part.blob.decode('utf-8')
            assert "w:comment" in blob_str
            assert "主角称号说明" in blob_str
            break
    assert comments_found is True
