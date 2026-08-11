import os
from fastapi import APIRouter
from fastapi.responses import FileResponse
from pydantic import BaseModel
from docx import Document as DocxDocument
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from app.core.document import EMPTY_SKIP_STYLES
from datetime import datetime
from app.core.database import (
    get_project, get_current_document, get_errors, get_error,
    update_error_status, update_error_suggested, update_project_status,
    get_paragraph_by_idx, get_paragraph_by_uuid, update_paragraph_revised, get_revised_paragraphs,
    get_chapters, get_annotations,
)

router = APIRouter()


def _apply_heading_style(docx, para, level: int):
    """安全写入 Word 标题样式，兼容中英文环境（Heading 1 与 标题 1）。"""
    targets = [f"Heading {level}", f"标题 {level}"]
    for t in targets:
        try:
            if t in docx.styles:
                para.style = docx.styles[t]
                return
        except Exception:
            pass
    for t in targets:
        try:
            para.style = t
            return
        except Exception:
            pass


def _recompute_paragraph(document_id: str, paragraph_idx_or_uuid: int | str):
    """根据该段所有已采纳错误，重算 revised_text（撤回即重新基于原文计算）。
    支持传入 paragraph_idx (int/str) 或 paragraph_uuid (str)。
    """
    if isinstance(paragraph_idx_or_uuid, str) and not paragraph_idx_or_uuid.isdigit():
        para = get_paragraph_by_uuid(document_id, paragraph_idx_or_uuid)
    else:
        para = get_paragraph_by_idx(document_id, int(paragraph_idx_or_uuid))
    if not para:
        return
    para_idx = para["idx"]
    para_uuid = para.get("uuid")
    accepted = [
        e for e in get_errors(document_id)
        if ((para_uuid and e.get("paragraph_uuid") == para_uuid) or e.get("paragraph_index") == para_idx)
        and e["user_status"] == "accepted"
    ]
    original_clean = para["text"].lstrip(" \t\r\n\u3000")
    has_lstrip = original_clean != para["text"]
    base_text = original_clean if has_lstrip else para["text"]

    if not accepted:
        if para.get("edit_note") and para.get("revised_text") and para["revised_text"] != para["text"]:
            return
        if has_lstrip:
            update_paragraph_revised(para["id"], original_clean)
        else:
            update_paragraph_revised(para["id"], None)
        return

    # 用 indexOf 获取每个错误在基底文本中的确切位置
    indexed: list[tuple[int, dict]] = []
    for e in accepted:
        pos = base_text.find(e["original_text"])
        if pos >= 0:
            indexed.append((pos, e))

    # 按位置降序处理（从右向左），保证左侧未处理的位置不受长度变化影响
    indexed.sort(key=lambda x: x[0], reverse=True)

    revised = base_text
    for pos, e in indexed:
        end = pos + len(e["original_text"])
        if revised[pos:end] == e["original_text"]:
            revised = revised[:pos] + e["suggested_text"] + revised[end:]

    update_paragraph_revised(para["id"], revised)


class StatusBody(BaseModel):
    status: str  # accepted | rejected | pending
    custom_text: str | None = None


@router.post("/projects/{project_id}/errors/{error_id}/status")
async def set_error_status(project_id: str, error_id: int, body: StatusBody):
    if body.status not in ("accepted", "rejected", "pending"):
        return {"error": "非法状态"}
    if body.custom_text and body.status == "accepted":
        update_error_suggested(error_id, body.custom_text)
    update_error_status(error_id, body.status)
    e = get_error(error_id)
    if e:
        _recompute_paragraph(e["document_id"], e.get("paragraph_uuid") or e["paragraph_index"])
    return {"status": "ok"}


@router.post("/projects/{project_id}/accept-all")
async def accept_all(project_id: str):
    doc = get_current_document(project_id)
    if not doc:
        return {"error": "项目无文档"}
    doc_id = doc["id"]
    errors = get_errors(doc_id)
    if not errors:
        return {"status": "ok", "count": 0}
    for e in errors:
        update_error_status(e["id"], "accepted")
    # 按段落分组，每段只重算一次（Bug 3 修复）
    seen = set()
    for e in errors:
        key = e.get("paragraph_uuid") or e["paragraph_index"]
        if key not in seen:
            seen.add(key)
            _recompute_paragraph(doc_id, key)
    return {"status": "ok", "count": len(errors)}


from docx.opc.packuri import PackURI
from docx.opc.part import XmlPart
from docx.opc.constants import RELATIONSHIP_TYPE


def _ensure_comments_part(docx):
    """查找或初始化 word/comments.xml OpenXML Part，并注册 content type 与 relation"""
    for rel in docx.part.rels.values():
        if "comments" in str(rel.target_ref):
            target = rel.target_part
            if hasattr(target, "element"):
                return target.element
            elif hasattr(target, "_element"):
                return target._element
            else:
                c_elm = parse_xml(target.blob)
                xml_part = XmlPart(
                    PackURI('/word/comments.xml'),
                    'application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml',
                    c_elm,
                    docx.part.package
                )
                rel._target = xml_part
                return c_elm

    comments_elm = parse_xml(
        r'<w:comments %s %s></w:comments>' % (
            nsdecls('w'), nsdecls('r')
        )
    )
    comments_part = XmlPart(
        PackURI('/word/comments.xml'),
        'application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml',
        comments_elm,
        docx.part.package
    )
    docx.part.relate_to(
        comments_part,
        RELATIONSHIP_TYPE.COMMENTS
    )
    return comments_elm


from docx.shared import Pt


@router.post("/projects/{project_id}/export")
async def export_document(project_id: str, export_mode: str = "print"):
    """导出校稿版 docx。

    export_mode:
      - "print": 打印/出版版，正文嵌入右上标 [注X] 角标，章末生成 【本章注释】 规范列表。
      - "comment": 批注版，直接向 docx 写入 Word 原生 w:comment 侧栏气泡批注。
    """
    doc = get_current_document(project_id)
    if not doc:
        return {"error": "项目无文档"}

    doc_id = doc["id"]
    file_path = doc.get("file_path") or ""
    if not file_path or not os.path.exists(file_path):
        return {"error": "原始文件不存在，无法导出"}

    paras = get_revised_paragraphs(doc_id)

    os.makedirs("backend/static/exports", exist_ok=True)

    proj = get_project(project_id)
    raw_name = (proj.get("name") if proj else None) or "校稿文档"
    clean_name = "".join(c for c in raw_name if c not in r'/\:*?"<>|').strip() or "校稿文档"
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    mode_tag = "批注版" if export_mode == "comment" else "打印版"
    fname = f"{clean_name}_{mode_tag}_{ts}.docx"

    if file_path and os.path.exists(file_path):
        docx = DocxDocument(file_path)
        body = docx._element.body

        removable_tags = {qn("w:p"), qn("w:tbl"), qn("w:sdt")}
        for child in list(body):
            if child.tag in removable_tags:
                body.remove(child)

        from app.core.database import get_project_style_config
        style_cfg = get_project_style_config(project_id)
        indent_enabled = style_cfg.get("first_line_indent_enabled", False)

        chapters = get_chapters(doc_id)
        ch_map = {}
        for c in chapters:
            if c.get("title_paragraph_uuid"):
                ch_map[c["title_paragraph_uuid"]] = c
            if c.get("title_paragraph_idx") is not None:
                ch_map[c["title_paragraph_idx"]] = c

        all_annots = get_annotations(doc_id)
        annots_by_para = {}
        for a in all_annots:
            p_uuid = a.get("paragraph_uuid")
            p_idx = a.get("paragraph_idx")
            if p_uuid:
                annots_by_para.setdefault(str(p_uuid), []).append(a)
            if p_idx is not None:
                annots_by_para.setdefault(str(p_idx), []).append(a)

        current_chapter_annots = []
        annot_counter = 0
        comment_id_counter = 0

        def _flush_chapter_annotations():
            nonlocal current_chapter_annots, annot_counter
            if not current_chapter_annots or export_mode != "print":
                return

            note_heading = docx.add_paragraph()
            pPr = note_heading._element.get_or_add_pPr()
            pPr.append(parse_xml(r'<w:spacing %s w:before="240" w:after="120"/>' % nsdecls("w")))

            r_h = note_heading.add_run("【本章注释】")
            r_h.bold = True
            r_h.font.size = Pt(10.5)

            for item in current_chapter_annots:
                num = item["num"]
                annot = item["annot"]
                p = docx.add_paragraph()
                p_pPr = p._element.get_or_add_pPr()
                p_pPr.append(parse_xml(r'<w:ind %s w:left="420" w:firstLine="-420"/>' % nsdecls("w")))
                p_pPr.append(parse_xml(r'<w:spacing %s w:before="0" w:after="60" w:line="280" w:lineRule="auto"/>' % nsdecls("w")))

                r_num = p.add_run(f"[{num}] ")
                r_num.bold = True
                r_num.font.size = Pt(9.0)

                r_sel = p.add_run(f"「{annot.get('selected_text', '')}」")
                r_sel.bold = True
                r_sel.font.size = Pt(9.0)

                r_cnt = p.add_run(f"：{annot.get('content', '')}")
                r_cnt.font.size = Pt(9.0)

            current_chapter_annots = []
            annot_counter = 0

        processed_annot_ids = set()

        for p_data in paras:
            text = p_data.get("text") or ""
            style_name = p_data.get("style_name") or "Normal"
            pb_type = p_data.get("page_break_type", "none")
            idx = p_data.get("idx", 0)
            p_uuid = p_data.get("uuid")

            ch_info = (ch_map.get(p_uuid) if p_uuid else None) or ch_map.get(idx)

            if ch_info and idx > 0 and export_mode == "print":
                _flush_chapter_annotations()

            new_para = docx.add_paragraph()

            if ch_info:
                level = ch_info.get("level", 1)
                _apply_heading_style(docx, new_para, level)
            else:
                try:
                    new_para.style = docx.styles[style_name]
                except (KeyError, Exception):
                    try:
                        new_para.style = docx.styles["Normal"]
                    except Exception:
                        pass

            para_annots = []
            raw_annots = (annots_by_para.get(str(p_uuid)) if p_uuid else None) or annots_by_para.get(str(idx)) or []
            for a in raw_annots:
                if a["id"] not in processed_annot_ids:
                    processed_annot_ids.add(a["id"])
                    para_annots.append(a)

            rendered_text = text
            if para_annots and rendered_text and export_mode == "print":
                curr_pos = 0
                for a in para_annots:
                    sel = a.get("selected_text", "")
                    pos = rendered_text.find(sel, curr_pos) if sel else -1
                    if pos >= 0:
                        new_para.add_run(rendered_text[curr_pos:pos + len(sel)])
                        annot_counter += 1
                        num_str = f"注{annot_counter}"

                        tag_run = new_para.add_run(f"[{num_str}]")
                        tag_run.font.superscript = True
                        tag_run.font.size = Pt(8.5)

                        current_chapter_annots.append({
                            "num": num_str,
                            "annot": a
                        })
                        curr_pos = pos + len(sel)

                if curr_pos < len(rendered_text):
                    new_para.add_run(rendered_text[curr_pos:])
            elif para_annots and rendered_text and export_mode == "comment":
                comments_elm = _ensure_comments_part(docx)
                curr_pos = 0
                for a in para_annots:
                    sel = a.get("selected_text", "")
                    content = a.get("content", "")
                    pos = rendered_text.find(sel, curr_pos) if sel else -1
                    if pos >= 0:
                        if pos > curr_pos:
                            new_para.add_run(rendered_text[curr_pos:pos])

                        comment_id_str = str(comment_id_counter)
                        comment_id_counter += 1

                        c_xml = parse_xml(
                            r'<w:comment %s w:id="%s" w:author="校稿助手" w:date="%s"><w:p><w:r><w:t>%s</w:t></w:r></w:p></w:comment>' % (
                                nsdecls('w'),
                                comment_id_str,
                                datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ"),
                                content.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                            )
                        )
                        comments_elm.append(c_xml)

                        start_node = parse_xml(r'<w:commentRangeStart %s w:id="%s"/>' % (nsdecls('w'), comment_id_str))
                        new_para._element.append(start_node)

                        new_para.add_run(sel)

                        end_node = parse_xml(r'<w:commentRangeEnd %s w:id="%s"/>' % (nsdecls('w'), comment_id_str))
                        new_para._element.append(end_node)

                        ref_node = parse_xml(r'<w:r %s><w:commentReference %s w:id="%s"/></w:r>' % (nsdecls('w'), nsdecls('w'), comment_id_str))
                        new_para._element.append(ref_node)

                        curr_pos = pos + len(sel)

                if curr_pos < len(rendered_text):
                    new_para.add_run(rendered_text[curr_pos:])
            else:
                new_para.add_run(rendered_text)

            if pb_type in ("original", "manual", "auto_chapter") and idx > 0:
                pPr = new_para._element.get_or_add_pPr()
                if pPr.find(qn("w:pageBreakBefore")) is None:
                    pPr.append(parse_xml(r'<w:pageBreakBefore %s/>' % nsdecls("w")))

            if indent_enabled and not ch_info and "heading" not in style_name.lower() and "标题" not in style_name:
                pPr = new_para._element.get_or_add_pPr()
                if pPr.find(qn("w:ind")) is None:
                    pPr.append(parse_xml(r'<w:ind %s w:firstLineChars="200" w:firstLine="420"/>' % nsdecls("w")))

        if export_mode == "print":
            _flush_chapter_annotations()

        fpath = os.path.join("backend/static/exports", fname)
        docx.save(fpath)

    update_project_status(project_id, "completed")
    return FileResponse(fpath, filename=fname)
