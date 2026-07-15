from fastapi import APIRouter
from fastapi.responses import StreamingResponse
from io import BytesIO
from docx import Document

router = APIRouter()


@router.get("/export/{document_id}")
async def export_document(document_id: str):
    """导出校对后的 docx 文件（骨架版本返回空白 docx）。"""
    doc = Document()
    doc.add_heading("校对结果", 0)
    doc.add_paragraph(f"文档 ID：{document_id}")
    doc.add_paragraph("（骨架版本：此为占位文件，后续将生成包含修订标记的完整文档）")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename=proofread_{document_id}.docx"
        },
    )
