import re
from docx import Document as DocxDocument
from docx.oxml.ns import qn

EMPTY_SKIP_STYLES = {"TOC", "Header", "Footer", "Footnote Text", "Endnote Text"}

CHAPTER_REGEX_L1 = re.compile(r"^第[零一二三四五六七八九十百千万0-9A-Za-z]+[章卷部]")
CHAPTER_REGEX_L2 = re.compile(r"^第[零一二三四五六七八九十百千万0-9A-Za-z]+[节回]")
CHAPTER_REGEX_NUMBER = re.compile(r"^([0-9]{1,3}|[一二三四五六七八九十]+)[\s.、:：_-]")
CHAPTER_KEYWORDS = {"序章", "序言", "前言", "楔子", "尾声", "后记", "附录", "番外"}


def extract_word_comments(doc: DocxDocument) -> list[dict]:
    """使用纯 OpenXML 节点遍历提取 Word docx 中的原生批注 (w:commentReference / w:comment)"""
    extracted_annotations = []
    try:
        comments_part = None
        for rel in doc.part.rels.values():
            if "comments" in str(rel.target_ref):
                comments_part = rel.target_part
                break

        if not comments_part:
            return extracted_annotations

        comments_xml = comments_part.element
        comments_dict = {}
        for comment_elem in comments_xml.xpath(".//w:comment"):
            c_id = comment_elem.get(qn("w:id")) or comment_elem.get("id")
            c_text = "".join(t.text for t in comment_elem.xpath(".//w:t") if t.text)
            if c_id and c_text:
                comments_dict[c_id] = c_text

        if not comments_dict:
            return extracted_annotations

        para_idx = 0
        for para in doc.paragraphs:
            style_name = para.style.name or ""
            if style_name in EMPTY_SKIP_STYLES:
                continue

            elem = para._element
            ref_elems = elem.xpath(".//w:commentReference")
            for ref in ref_elems:
                c_id = ref.get(qn("w:id")) or ref.get("id")
                if c_id in comments_dict:
                    nodes = elem.xpath(f".//w:commentRangeStart[@w:id='{c_id}']/following-sibling::w:r[following-sibling::w:commentRangeEnd[@w:id='{c_id}']]")
                    selected_text = "".join("".join(r.xpath(".//w:t/text()")) for r in nodes).strip()
                    if not selected_text:
                        selected_text = (para.text or "")[:20]

                    extracted_annotations.append({
                        "paragraph_idx": para_idx,
                        "selected_text": selected_text or "批注引用",
                        "content": comments_dict[c_id],
                    })
            para_idx += 1
    except Exception:
        pass

    return extracted_annotations


def parse_paragraphs(file_path: str) -> tuple[list[tuple], list[dict], bool, list[dict]]:
    """解析 docx，返回:
    1. 有序段落列表 [(idx, text, style_name, page_break_type), ...]（包含空段落）。
    2. 基于标题样式和正则初步提取的初始章节列表。
    3. 智能检测是否开启首行缩进 (has_first_line_indent)。
    4. 从 Word 原生批注中提取的注释列表。
    """
    doc = DocxDocument(file_path)
    extracted_annotations = extract_word_comments(doc)

    rows = []
    chapters = []
    idx = 0

    body_para_count = 0
    indented_para_count = 0

    for para in doc.paragraphs:
        style_name = para.style.name or ""
        if style_name in EMPTY_SKIP_STYLES:
            continue

        raw_text = para.text or ""
        text = raw_text.rstrip("\r\n")

        # 1. 精准检测物理硬分页符类型（排查软分页与 w:sectPr 页面边距/分节符）
        elem = para._element
        has_original_break = False
        if hasattr(elem, "pPr") and elem.pPr is not None:
            if elem.pPr.find(qn("w:pageBreakBefore")) is not None:
                has_original_break = True
        if not has_original_break:
            for br in elem.xpath(".//w:br"):
                br_type = br.get(qn("w:type")) or br.get("type")
                if br_type == "page":
                    has_original_break = True
                    break

        # 2. 识别章节级别
        chapter_level = None
        style_lower = style_name.lower()

        if "heading 1" in style_lower or "标题 1" in style_name or style_lower == "heading1" or style_name == "标题1":
            chapter_level = 1
        elif "heading 2" in style_lower or "标题 2" in style_name or style_lower == "heading2" or style_name == "标题2":
            chapter_level = 2
        elif "heading 3" in style_lower or "标题 3" in style_name or style_lower == "heading3" or style_name == "标题3":
            chapter_level = 3
        elif "heading" in style_lower or "标题" in style_name or "title" in style_lower:
            chapter_level = 1
        elif text.strip():
            t_strip = text.strip()
            if CHAPTER_REGEX_L1.match(t_strip) or t_strip in CHAPTER_KEYWORDS:
                chapter_level = 1
            elif CHAPTER_REGEX_L2.match(t_strip):
                chapter_level = 2
            elif CHAPTER_REGEX_NUMBER.match(t_strip) and len(t_strip) < 40:
                chapter_level = 2

        # 3. 智能检测正文段落首行缩进特征 (排除标题与纯空段)
        if chapter_level is None and text.strip():
            body_para_count += 1
            is_indented = False

            # a) Word XML 格式层检测 (<w:ind w:firstLine="..." /> 或 <w:ind w:firstLineChars="..." />)
            pPr = elem.find(qn("w:pPr"))
            if pPr is not None:
                ind = pPr.find(qn("w:ind"))
                if ind is not None:
                    first_line = ind.get(qn("w:firstLine")) or ind.get("firstLine")
                    first_line_chars = ind.get(qn("w:firstLineChars")) or ind.get("firstLineChars")
                    if (first_line and first_line != "0") or (first_line_chars and first_line_chars != "0"):
                        is_indented = True

            # b) Python-docx 对象的 paragraph_format
            if not is_indented:
                try:
                    fl_indent = para.paragraph_format.first_line_indent
                    if fl_indent is not None and fl_indent > 0:
                        is_indented = True
                except Exception:
                    pass

            # c) 文本前导硬空格/Tab检测 (全角空格 \u3000, 制表符 \t, 2+ 连续半角空格)
            if not is_indented:
                if raw_text.startswith("\u3000") or raw_text.startswith("\t") or raw_text.startswith("  "):
                    is_indented = True

            if is_indented:
                indented_para_count += 1

        # 4. 确定最终 Page Break 属性 ('original' | 'auto_chapter' | 'none')
        page_break_type = "none"
        if has_original_break and chapter_level == 1 and idx > 0:
            page_break_type = "auto_chapter"
        elif has_original_break:
            page_break_type = "original"
        elif chapter_level == 1 and idx > 0:
            page_break_type = "auto_chapter"

        rows.append((idx, text, style_name, page_break_type))

        if chapter_level is not None and text.strip():
            chapters.append({
                "title": text.strip(),
                "title_paragraph_idx": idx,
                "level": chapter_level,
                "detected_by": "original",
            })

        idx += 1

    total_paras = len(rows)
    for i, ch in enumerate(chapters):
        ch["start_idx"] = ch["title_paragraph_idx"]
        if i + 1 < len(chapters):
            ch["end_idx"] = max(ch["start_idx"], chapters[i + 1]["title_paragraph_idx"] - 1)
        else:
            ch["end_idx"] = max(ch["start_idx"], total_paras - 1 if total_paras > 0 else 0)

    has_first_line_indent = False
    if body_para_count > 0:
        ratio = indented_para_count / body_para_count
        if ratio >= 0.5 or (body_para_count <= 5 and indented_para_count >= 2):
            has_first_line_indent = True

    return rows, chapters, has_first_line_indent, extracted_annotations
